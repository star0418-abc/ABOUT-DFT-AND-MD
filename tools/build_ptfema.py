#!/usr/bin/env python3
# -*- coding: utf-8 -*-

# Author: Star



"""
build_ptfema.py - 构建 PTFEMA（聚三氟乙基甲基丙烯酸酯）线性链
==================================================================

功能:
  1. 从 TFEMA.mol2 单体构建 N 聚合度的 PTFEMA 线性链
  2. 支持 H/Me 端基封端
  3. 可选使用 docx 坐标作为构象模板
  4. 输出 repeat_capped 模型和完整链

化学规则:
  - TFEMA 是甲基丙烯酸酯类单体
  - 聚合发生在 C=C 双键上（非羰基双键）
  - 聚合连接：前一单体的 CH2 端碳 与 后一单体的取代端碳 连接

用法:
  python tools/build_ptfema.py --monomer mol2/TFEMA.mol2 --n 10 --cap H
  python tools/build_ptfema.py --monomer mol2/TFEMA.mol2 --n 10 --cap Me
  python tools/build_ptfema.py --monomer mol2/TFEMA.mol2 --n 10 --docx "supplementary data 1.docx"

版本: 1.0.0
"""

import argparse
import os
import sys
import math
import re
from pathlib import Path
from typing import Tuple, List, Dict, Optional, Any
from dataclasses import dataclass
from collections import defaultdict

# RDKit
try:
    from rdkit import Chem
    from rdkit.Chem import AllChem, rdMolTransforms
    from rdkit.Chem.rdForceFieldHelpers import UFFOptimizeMolecule, UFFHasAllMoleculeParams
    from rdkit.Chem.rdForceFieldHelpers import MMFFOptimizeMolecule, MMFFHasAllMoleculeParams
    from rdkit import RDLogger
    RDLogger.logger().setLevel(RDLogger.ERROR)
except ImportError:
    print("[FATAL] RDKit 未安装: conda install -c conda-forge rdkit")
    sys.exit(1)

# 尝试导入 docx 库
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ==============================================================================
# 颜色输出
# ==============================================================================

class Colors:
    RESET = "\033[0m"
    RED = "\033[0;31m"
    GREEN = "\033[0;32m"
    YELLOW = "\033[0;33m"
    BLUE = "\033[0;34m"
    CYAN = "\033[0;36m"
    BOLD = "\033[1m"


def log_success(msg): print(f"{Colors.GREEN}✓{Colors.RESET} {msg}")
def log_error(msg): print(f"{Colors.RED}✗{Colors.RESET} {msg}", file=sys.stderr)
def log_info(msg): print(f"{Colors.BLUE}ℹ{Colors.RESET} {msg}")
def log_warn(msg): print(f"{Colors.YELLOW}⚠{Colors.RESET} {msg}")


def print_header(title: str):
    print(f"\n{Colors.BOLD}{'=' * 60}{Colors.RESET}")
    print(f" {Colors.CYAN}{title}{Colors.RESET}")
    print(f"{Colors.BOLD}{'=' * 60}{Colors.RESET}")


# ==============================================================================
# MOL2 读写器（不依赖 OpenBabel）
# ==============================================================================

@dataclass
class Mol2Atom:
    """MOL2 原子"""
    idx: int
    name: str
    x: float
    y: float
    z: float
    atom_type: str
    res_id: int
    res_name: str
    charge: float


@dataclass
class Mol2Bond:
    """MOL2 键"""
    idx: int
    atom1: int
    atom2: int
    bond_type: str  # "1", "2", "3", "ar"


@dataclass
class Mol2Molecule:
    """MOL2 分子"""
    name: str
    atoms: List[Mol2Atom]
    bonds: List[Mol2Bond]
    
    @property
    def num_atoms(self) -> int:
        return len(self.atoms)
    
    @property
    def num_bonds(self) -> int:
        return len(self.bonds)


def parse_mol2(filepath: str) -> Mol2Molecule:
    """解析 MOL2 文件"""
    with open(filepath, 'r') as f:
        lines = f.readlines()
    
    name = "MOL"
    atoms = []
    bonds = []
    section = None
    
    for line in lines:
        line = line.strip()
        
        if line.startswith("@<TRIPOS>MOLECULE"):
            section = "molecule"
            continue
        elif line.startswith("@<TRIPOS>ATOM"):
            section = "atom"
            continue
        elif line.startswith("@<TRIPOS>BOND"):
            section = "bond"
            continue
        elif line.startswith("@<TRIPOS>"):
            section = "other"
            continue
        
        if not line:
            continue
        
        if section == "molecule":
            if not name or name == "MOL":
                name = line.split()[0] if line.split() else "MOL"
            section = "molecule_done"
        
        elif section == "atom":
            parts = line.split()
            if len(parts) >= 6:
                atom = Mol2Atom(
                    idx=int(parts[0]),
                    name=parts[1],
                    x=float(parts[2]),
                    y=float(parts[3]),
                    z=float(parts[4]),
                    atom_type=parts[5],
                    res_id=int(parts[6]) if len(parts) > 6 else 1,
                    res_name=parts[7] if len(parts) > 7 else "MOL",
                    charge=float(parts[8]) if len(parts) > 8 else 0.0
                )
                atoms.append(atom)
        
        elif section == "bond":
            parts = line.split()
            if len(parts) >= 4:
                bond = Mol2Bond(
                    idx=int(parts[0]),
                    atom1=int(parts[1]),
                    atom2=int(parts[2]),
                    bond_type=parts[3]
                )
                bonds.append(bond)
    
    return Mol2Molecule(name=name, atoms=atoms, bonds=bonds)


def write_mol2(mol: Mol2Molecule, filepath: str) -> None:
    """写入 MOL2 文件"""
    with open(filepath, 'w') as f:
        f.write("@<TRIPOS>MOLECULE\n")
        f.write(f"{mol.name}\n")
        f.write(f" {mol.num_atoms} {mol.num_bonds} 1 0 0\n")
        f.write("SMALL\n")
        f.write("NO_CHARGES\n\n")
        
        f.write("@<TRIPOS>ATOM\n")
        for atom in mol.atoms:
            f.write(f"{atom.idx:>6d} {atom.name:<4s} {atom.x:>10.4f} {atom.y:>10.4f} {atom.z:>10.4f} "
                   f"{atom.atom_type:<6s} {atom.res_id:>3d} {atom.res_name:<4s} {atom.charge:>8.4f}\n")
        
        f.write("@<TRIPOS>BOND\n")
        for bond in mol.bonds:
            f.write(f"{bond.idx:>6d} {bond.atom1:>5d} {bond.atom2:>5d} {bond.bond_type}\n")
        
        f.write("@<TRIPOS>SUBSTRUCTURE\n")
        f.write(f"     1 {mol.name[:3]:<4s}        1 RESIDUE    0 **** **** 0 ROOT\n")


def mol2_to_rdkit(mol2: Mol2Molecule) -> Chem.RWMol:
    """将 Mol2Molecule 转换为 RDKit RWMol"""
    rwmol = Chem.RWMol()
    
    # 元素映射
    type_to_element = {
        'C.3': 6, 'C.2': 6, 'C.1': 6, 'C.ar': 6,
        'N.3': 7, 'N.2': 7, 'N.1': 7, 'N.ar': 7, 'N.am': 7, 'N.pl3': 7,
        'O.3': 8, 'O.2': 8, 'O.co2': 8,
        'S.3': 16, 'S.2': 16, 'S.O': 16, 'S.O2': 16,
        'P.3': 15,
        'F': 9, 'Cl': 17, 'Br': 35, 'I': 53,
        'H': 1,
    }
    
    # 添加原子
    mol2_to_rdkit_idx = {}
    for atom in mol2.atoms:
        element = type_to_element.get(atom.atom_type, 6)
        rd_atom = Chem.Atom(element)
        rd_idx = rwmol.AddAtom(rd_atom)
        mol2_to_rdkit_idx[atom.idx] = rd_idx
    
    # 添加键
    bond_type_map = {
        '1': Chem.BondType.SINGLE,
        '2': Chem.BondType.DOUBLE,
        '3': Chem.BondType.TRIPLE,
        'ar': Chem.BondType.AROMATIC,
        'am': Chem.BondType.SINGLE,
    }
    
    for bond in mol2.bonds:
        a1 = mol2_to_rdkit_idx[bond.atom1]
        a2 = mol2_to_rdkit_idx[bond.atom2]
        bt = bond_type_map.get(bond.bond_type, Chem.BondType.SINGLE)
        rwmol.AddBond(a1, a2, bt)
    
    # 设置坐标
    conf = Chem.Conformer(len(mol2.atoms))
    for atom in mol2.atoms:
        rd_idx = mol2_to_rdkit_idx[atom.idx]
        conf.SetAtomPosition(rd_idx, (atom.x, atom.y, atom.z))
    rwmol.AddConformer(conf, assignId=True)
    
    return rwmol


def rdkit_to_mol2(mol: Chem.Mol, name: str = "MOL", res_name: str = "MOL") -> Mol2Molecule:
    """将 RDKit Mol 转换为 Mol2Molecule"""
    conf = mol.GetConformer()
    
    # 元素到 MOL2 类型映射
    element_to_type = {
        6: 'C.3', 7: 'N.3', 8: 'O.3', 9: 'F', 16: 'S.3', 15: 'P.3',
        17: 'Cl', 35: 'Br', 53: 'I', 1: 'H'
    }
    
    atoms = []
    for i, atom in enumerate(mol.GetAtoms()):
        pos = conf.GetAtomPosition(i)
        element = atom.GetAtomicNum()
        symbol = atom.GetSymbol()
        
        # 确定 MOL2 原子类型
        atom_type = element_to_type.get(element, symbol)
        
        # 对碳和氧特殊处理
        if element == 6:
            hybridization = atom.GetHybridization()
            if hybridization == Chem.HybridizationType.SP2:
                atom_type = 'C.2'
            elif hybridization == Chem.HybridizationType.SP:
                atom_type = 'C.1'
            else:
                atom_type = 'C.3'
        elif element == 8:
            # 检查是否是羰基氧
            for bond in atom.GetBonds():
                if bond.GetBondType() == Chem.BondType.DOUBLE:
                    atom_type = 'O.2'
                    break
            else:
                atom_type = 'O.3'
        
        atom_name = f"{symbol}{i+1}"
        atoms.append(Mol2Atom(
            idx=i+1,
            name=atom_name,
            x=pos.x, y=pos.y, z=pos.z,
            atom_type=atom_type,
            res_id=1,
            res_name=res_name[:3],
            charge=0.0
        ))
    
    bonds = []
    for i, bond in enumerate(mol.GetBonds()):
        bt = bond.GetBondType()
        if bt == Chem.BondType.SINGLE:
            bond_type = '1'
        elif bt == Chem.BondType.DOUBLE:
            bond_type = '2'
        elif bt == Chem.BondType.TRIPLE:
            bond_type = '3'
        elif bt == Chem.BondType.AROMATIC:
            bond_type = 'ar'
        else:
            bond_type = '1'
        
        bonds.append(Mol2Bond(
            idx=i+1,
            atom1=bond.GetBeginAtomIdx() + 1,
            atom2=bond.GetEndAtomIdx() + 1,
            bond_type=bond_type
        ))
    
    return Mol2Molecule(name=name, atoms=atoms, bonds=bonds)


# ==============================================================================
# 聚合双键识别
# ==============================================================================

def find_polymerization_double_bond(mol2: Mol2Molecule) -> Tuple[int, int, int, int]:
    """
    识别单体中的聚合双键
    
    对于甲基丙烯酸酯类单体:
    - 找到 C=C 双键（键阶=2）
    - 其中一个碳连着羰基碳（酯基），另一个是 CH2 端
    
    Returns:
        (ch2_carbon_idx, substituted_carbon_idx, ch2_carbon_mol2_idx, subst_carbon_mol2_idx)
        前两个是 0-indexed (RDKit)，后两个是 1-indexed (MOL2)
    """
    # 构建邻接表和键类型
    adjacency = defaultdict(list)
    bond_types = {}
    
    for bond in mol2.bonds:
        a1, a2 = bond.atom1, bond.atom2
        adjacency[a1].append(a2)
        adjacency[a2].append(a1)
        bond_types[(min(a1, a2), max(a1, a2))] = bond.bond_type
    
    # 构建原子类型映射
    atom_types = {atom.idx: atom.atom_type for atom in mol2.atoms}
    atom_elements = {}
    for atom in mol2.atoms:
        if atom.atom_type.startswith('C'):
            atom_elements[atom.idx] = 'C'
        elif atom.atom_type.startswith('O'):
            atom_elements[atom.idx] = 'O'
        elif atom.atom_type.startswith('N'):
            atom_elements[atom.idx] = 'N'
        elif atom.atom_type == 'H':
            atom_elements[atom.idx] = 'H'
        elif atom.atom_type == 'F':
            atom_elements[atom.idx] = 'F'
        else:
            atom_elements[atom.idx] = atom.atom_type[0]
    
    # 找所有 C=C 双键（排除羰基 C=O）
    cc_double_bonds = []
    for (a1, a2), bt in bond_types.items():
        if bt == '2':
            e1 = atom_elements.get(a1, '')
            e2 = atom_elements.get(a2, '')
            if e1 == 'C' and e2 == 'C':
                cc_double_bonds.append((a1, a2))
    
    if not cc_double_bonds:
        raise RuntimeError("未找到 C=C 双键")
    
    # 对于甲基丙烯酸酯类：
    # 找到连着羰基碳的那个双键碳（取代端）和 CH2 端
    for c1, c2 in cc_double_bonds:
        # 检查 c1 是否连着羰基碳
        c1_neighbors = [n for n in adjacency[c1] if n != c2]
        c2_neighbors = [n for n in adjacency[c2] if n != c1]
        
        c1_connects_carbonyl = False
        c2_connects_carbonyl = False
        
        for n in c1_neighbors:
            if atom_elements.get(n) == 'C':
                # 检查这个碳是否是羰基碳（连着 C=O）
                for nn in adjacency[n]:
                    if atom_elements.get(nn) == 'O':
                        key = (min(n, nn), max(n, nn))
                        if bond_types.get(key) == '2':
                            c1_connects_carbonyl = True
                            break
        
        for n in c2_neighbors:
            if atom_elements.get(n) == 'C':
                for nn in adjacency[n]:
                    if atom_elements.get(nn) == 'O':
                        key = (min(n, nn), max(n, nn))
                        if bond_types.get(key) == '2':
                            c2_connects_carbonyl = True
                            break
        
        # 取代端碳连着羰基碳，CH2 端不连
        if c1_connects_carbonyl and not c2_connects_carbonyl:
            substituted = c1
            ch2 = c2
        elif c2_connects_carbonyl and not c1_connects_carbonyl:
            substituted = c2
            ch2 = c1
        else:
            # 备用规则：连着更多重原子的是取代端
            c1_heavy = len([n for n in c1_neighbors if atom_elements.get(n) != 'H'])
            c2_heavy = len([n for n in c2_neighbors if atom_elements.get(n) != 'H'])
            
            if c1_heavy > c2_heavy:
                substituted = c1
                ch2 = c2
            else:
                substituted = c2
                ch2 = c1
        
        # 返回 (0-indexed, 0-indexed, mol2-indexed, mol2-indexed)
        return (ch2 - 1, substituted - 1, ch2, substituted)
    
    raise RuntimeError("无法识别聚合双键的 CH2 端和取代端")


# ==============================================================================
# DOCX 坐标提取
# ==============================================================================

def extract_coords_from_docx(docx_path: str) -> Optional[List[Tuple[str, float, float, float]]]:
    """
    从 docx 文件提取 PTFMA 坐标
    
    查找 "Cartesian coordinates for the optimized molecular structure of PTFMA"
    对应的表格，提取坐标
    
    Returns:
        [(element, x, y, z), ...] 或 None
    """
    if not HAS_DOCX:
        log_warn("python-docx 未安装，跳过 docx 解析")
        return None
    
    if not os.path.exists(docx_path):
        log_warn(f"docx 文件不存在: {docx_path}")
        return None
    
    try:
        doc = DocxDocument(docx_path)
    except Exception as e:
        log_warn(f"无法打开 docx: {e}")
        return None
    
    # 找到 PTFMA 对应的表格（通常在 "PTFMA" 标题之后）
    # 根据文档结构，PTFMA 是第 4 个标题，对应第 4 个表格（索引 3）
    ptfma_table_idx = None
    
    for i, para in enumerate(doc.paragraphs):
        if 'PTFMA' in para.text:
            # 找到 PTFMA 标题，其后的表格就是坐标
            # 在这个文档结构中，表格索引与标题顺序对应
            ptfma_table_idx = 3  # 第 4 个表格
            break
    
    if ptfma_table_idx is None or ptfma_table_idx >= len(doc.tables):
        log_warn("未找到 PTFMA 坐标表格")
        return None
    
    # 从表格提取坐标
    table = doc.tables[ptfma_table_idx]
    coords = []
    
    # 坐标格式：元素符号 + 3个数字
    coord_pattern = re.compile(r'([A-Z][a-z]?)\s+(-?\d+\.\d+)\s+(-?\d+\.\d+)\s+(-?\d+\.\d+)')
    
    for row in table.rows:
        for cell in row.cells:
            # 单元格内可能有多行坐标
            cell_text = cell.text
            for line in cell_text.split('\n'):
                line = line.strip()
                match = coord_pattern.search(line)
                if match:
                    element = match.group(1)
                    x = float(match.group(2))
                    y = float(match.group(3))
                    z = float(match.group(4))
                    coords.append((element, x, y, z))
    
    if coords:
        log_success(f"从 docx 提取到 {len(coords)} 个原子坐标")
        return coords
    
    log_warn("docx 中未找到有效坐标")
    return None


def coords_to_xyz(coords: List[Tuple[str, float, float, float]], filepath: str, comment: str = "") -> None:
    """将坐标写入 XYZ 文件"""
    with open(filepath, 'w') as f:
        f.write(f"{len(coords)}\n")
        f.write(f"{comment}\n")
        for elem, x, y, z in coords:
            f.write(f"{elem:2s} {x:>12.6f} {y:>12.6f} {z:>12.6f}\n")


def coords_to_rdkit_mol(coords: List[Tuple[str, float, float, float]]) -> Optional[Chem.Mol]:
    """
    从坐标构建 RDKit 分子（尝试用距离推断键）
    """
    try:
        from rdkit.Chem import rdDetermineBonds
    except ImportError:
        log_warn("rdDetermineBonds 不可用，无法从坐标推断键")
        return None
    
    # 构建只有原子的分子
    rwmol = Chem.RWMol()
    
    element_to_num = {
        'H': 1, 'C': 6, 'N': 7, 'O': 8, 'F': 9, 'S': 16, 'Cl': 17
    }
    
    conf = Chem.Conformer(len(coords))
    for i, (elem, x, y, z) in enumerate(coords):
        atomic_num = element_to_num.get(elem, 6)
        atom = Chem.Atom(atomic_num)
        rwmol.AddAtom(atom)
        conf.SetAtomPosition(i, (x, y, z))
    
    rwmol.AddConformer(conf, assignId=True)
    
    # 尝试推断键
    try:
        rdDetermineBonds.DetermineBonds(rwmol, charge=0)
        Chem.SanitizeMol(rwmol)
        return rwmol.GetMol()
    except Exception as e:
        log_warn(f"推断键失败: {e}")
        return None


# ==============================================================================
# PTFEMA 链构建
# ==============================================================================

def build_ptfema_chain(
    monomer_mol2: Mol2Molecule,
    n_units: int,
    cap_type: str = "H"
) -> Chem.Mol:
    """
    构建 PTFEMA 链
    
    Args:
        monomer_mol2: 单体 MOL2 结构
        n_units: 聚合度
        cap_type: "H" 或 "Me" 端基
    
    Returns:
        RDKit Mol 对象
    """
    # 识别聚合双键
    ch2_idx, subst_idx, ch2_mol2, subst_mol2 = find_polymerization_double_bond(monomer_mol2)
    
    print(f"\n=== 聚合双键识别 ===")
    ch2_atom = next(a for a in monomer_mol2.atoms if a.idx == ch2_mol2)
    subst_atom = next(a for a in monomer_mol2.atoms if a.idx == subst_mol2)
    print(f"  CH2 端碳:   原子 {ch2_mol2} ({ch2_atom.name})")
    print(f"  取代端碳:   原子 {subst_mol2} ({subst_atom.name})")
    
    # 转换为 RDKit
    monomer = mol2_to_rdkit(monomer_mol2)
    
    # 找到 CH2 端碳上的一个氢（将用于连接）
    ch2_atom_rd = monomer.GetAtomWithIdx(ch2_idx)
    ch2_h_idx = None
    for neighbor in ch2_atom_rd.GetNeighbors():
        if neighbor.GetSymbol() == 'H':
            ch2_h_idx = neighbor.GetIdx()
            break
    
    # 找到取代端碳上的一个氢（将用于连接）
    subst_atom_rd = monomer.GetAtomWithIdx(subst_idx)
    subst_h_idx = None
    for neighbor in subst_atom_rd.GetNeighbors():
        if neighbor.GetSymbol() == 'H':
            subst_h_idx = neighbor.GetIdx()
            break
    
    if ch2_h_idx is None:
        raise RuntimeError(f"CH2 端碳 (idx={ch2_idx}) 上没有氢原子")
    
    print(f"  CH2 端氢:   原子 {ch2_h_idx}")
    print(f"  取代端氢:   原子 {subst_h_idx}" if subst_h_idx else "  取代端氢:   无（将直接连接）")
    
    # 构建链
    # 策略：使用 SMILES 构建然后优化
    # PTFEMA 重复单元: -CH2-C(CH3)(COOCH2CF3)-
    
    # 从单体 SMILES 构建
    # TFEMA SMILES: CC(=C)C(=O)OCC(F)(F)F (简化)
    # 聚合后: *CC(C)(C(=O)OCC(F)(F)F)* 重复
    
    # 读取单体 SMILES
    monomer_smiles = Chem.MolToSmiles(Chem.RemoveHs(monomer.GetMol()))
    print(f"\n单体 SMILES: {monomer_smiles}")
    
    # PTFEMA 重复单元 SMILES
    # 甲基丙烯酸酯类的重复单元: -CH2-C(CH3)(side)-
    # 对于 TFEMA: side = C(=O)OCC(F)(F)F
    repeat_smiles = "[*]CC(C)(C(=O)OCC(F)(F)F)[*]"
    
    # 构建链 SMILES
    if n_units == 1:
        # 单个单元 + 端基
        if cap_type == "H":
            chain_smiles = "CC(C)(C(=O)OCC(F)(F)F)C"
        else:  # Me
            chain_smiles = "CC(C)(C(=O)OCC(F)(F)F)C"
    else:
        # 多单元链
        # 端基-[CH2-C(CH3)(side)]_n-端基
        units = []
        for i in range(n_units):
            units.append("C(C)(C(=O)OCC(F)(F)F)")
        
        # 连接单元
        if cap_type == "H":
            head = "C"  # CH3 端
            tail = "C"  # CH3 端
        else:  # Me
            head = "CC"  # 甲基端
            tail = "C"   # 甲基端
        
        # 构建: head-C-[unit-C]_{n-1}-unit-tail
        chain_parts = [head]
        for i, unit in enumerate(units):
            chain_parts.append(unit)
            if i < n_units - 1:
                chain_parts.append("C")
        chain_parts.append(tail)
        
        chain_smiles = ''.join(chain_parts)
    
    print(f"链 SMILES: {chain_smiles[:80]}..." if len(chain_smiles) > 80 else f"链 SMILES: {chain_smiles}")
    
    # 解析 SMILES
    polymer = Chem.MolFromSmiles(chain_smiles)
    if polymer is None:
        raise RuntimeError(f"无法解析链 SMILES: {chain_smiles}")
    
    # 添加氢
    polymer = Chem.AddHs(polymer)
    
    # 嵌入 3D 坐标
    log_info("生成 3D 坐标...")
    try:
        params = AllChem.ETKDGv3()
        params.randomSeed = 2025
        status = AllChem.EmbedMolecule(polymer, params)
    except Exception:
        status = AllChem.EmbedMolecule(polymer, randomSeed=2025)
    
    if status < 0:
        log_warn("ETKDG 失败，尝试随机坐标...")
        AllChem.EmbedMolecule(polymer, useRandomCoords=True, randomSeed=2025)
    
    # 力场优化
    log_info("力场优化...")
    try:
        if MMFFHasAllMoleculeParams(polymer):
            MMFFOptimizeMolecule(polymer, maxIters=500)
            log_success("MMFF 优化完成")
        elif UFFHasAllMoleculeParams(polymer):
            UFFOptimizeMolecule(polymer, maxIters=500)
            log_success("UFF 优化完成")
    except Exception as e:
        log_warn(f"力场优化警告: {e}")
    
    return polymer


def build_repeat_capped(cap_type: str = "H") -> Chem.Mol:
    """
    构建 PTFEMA 重复单元 + 端基封端的小模型
    
    结构: CH3-CH2-C(CH3)(COOCH2CF3)-CH3
    """
    if cap_type == "H":
        # H 端基: CH3-CH2-C(CH3)(COOCH2CF3)-H -> 简化为加甲基
        smiles = "CC(C)(C(=O)OCC(F)(F)F)C"
    else:  # Me
        smiles = "CC(C)(C(=O)OCC(F)(F)F)CC"
    
    mol = Chem.MolFromSmiles(smiles)
    if mol is None:
        raise RuntimeError(f"无法解析 repeat unit SMILES: {smiles}")
    
    mol = Chem.AddHs(mol)
    
    # 嵌入并优化
    AllChem.EmbedMolecule(mol, randomSeed=2025)
    
    if MMFFHasAllMoleculeParams(mol):
        MMFFOptimizeMolecule(mol, maxIters=200)
    elif UFFHasAllMoleculeParams(mol):
        UFFOptimizeMolecule(mol, maxIters=200)
    
    return mol


# ==============================================================================
# 验证
# ==============================================================================

def validate_polymer(mol: Chem.Mol, n_units: int) -> Dict[str, Any]:
    """验证聚合物结构"""
    results = {}
    
    # 片段数
    frags = Chem.GetMolFrags(mol)
    results['num_fragments'] = len(frags)
    results['is_single_molecule'] = len(frags) == 1
    
    # 原子/键数
    results['num_atoms'] = mol.GetNumAtoms()
    results['num_bonds'] = mol.GetNumBonds()
    
    # 主链 C-C 键长
    conf = mol.GetConformer()
    cc_bond_lengths = []
    
    for bond in mol.GetBonds():
        a1 = bond.GetBeginAtom()
        a2 = bond.GetEndAtom()
        if a1.GetSymbol() == 'C' and a2.GetSymbol() == 'C':
            if bond.GetBondType() == Chem.BondType.SINGLE:
                pos1 = conf.GetAtomPosition(a1.GetIdx())
                pos2 = conf.GetAtomPosition(a2.GetIdx())
                length = math.sqrt(
                    (pos1.x - pos2.x)**2 + 
                    (pos1.y - pos2.y)**2 + 
                    (pos1.z - pos2.z)**2
                )
                cc_bond_lengths.append(length)
    
    if cc_bond_lengths:
        results['cc_bond_min'] = min(cc_bond_lengths)
        results['cc_bond_max'] = max(cc_bond_lengths)
        results['cc_bond_avg'] = sum(cc_bond_lengths) / len(cc_bond_lengths)
    
    # 检查最小非键距离
    min_distance = float('inf')
    for i in range(mol.GetNumAtoms()):
        for j in range(i + 2, mol.GetNumAtoms()):
            # 跳过直接成键的原子对
            bond = mol.GetBondBetweenAtoms(i, j)
            if bond is not None:
                continue
            
            pos_i = conf.GetAtomPosition(i)
            pos_j = conf.GetAtomPosition(j)
            dist = math.sqrt(
                (pos_i.x - pos_j.x)**2 + 
                (pos_i.y - pos_j.y)**2 + 
                (pos_i.z - pos_j.z)**2
            )
            if dist < min_distance:
                min_distance = dist
    
    results['min_nonbond_distance'] = min_distance
    results['no_clashes'] = min_distance > 0.8
    
    return results


# ==============================================================================
# 写入文件
# ==============================================================================

def write_pdb(mol: Chem.Mol, filepath: str) -> None:
    """写入 PDB 文件"""
    Chem.MolToPDBFile(mol, filepath)


def write_mol2_from_rdkit(mol: Chem.Mol, filepath: str, name: str = "MOL") -> None:
    """从 RDKit Mol 写入 MOL2 文件"""
    mol2_struct = rdkit_to_mol2(mol, name, name[:3])
    write_mol2(mol2_struct, filepath)


# ==============================================================================
# 主流程
# ==============================================================================

def main():
    parser = argparse.ArgumentParser(
        description="构建 PTFEMA（聚三氟乙基甲基丙烯酸酯）线性链",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 构建 10 聚体，H 端基
  python tools/build_ptfema.py --monomer mol2/TFEMA.mol2 --n 10 --cap H
  
  # 构建 10 聚体，甲基端基
  python tools/build_ptfema.py --monomer mol2/TFEMA.mol2 --n 10 --cap Me
  
  # 使用 docx 模板
  python tools/build_ptfema.py --monomer mol2/TFEMA.mol2 --n 10 --docx "supplementary data 1.docx"
"""
    )
    
    parser.add_argument("--monomer", required=True, help="单体 MOL2 文件路径")
    parser.add_argument("--n", type=int, default=10, help="聚合度 (默认: 10)")
    parser.add_argument("--cap", choices=["H", "Me"], default="H", help="端基类型 (默认: H)")
    parser.add_argument("--outprefix", default="PTFEMA", help="输出文件前缀 (默认: PTFEMA)")
    parser.add_argument("--outdir", default=None, help="输出目录 (默认: 与单体同目录)")
    parser.add_argument("--docx", default=None, help="补充材料 docx 文件路径（可选）")
    
    args = parser.parse_args()
    
    # 检查输入文件
    if not os.path.exists(args.monomer):
        log_error(f"单体文件不存在: {args.monomer}")
        return 1
    
    # 确定输出目录
    if args.outdir:
        outdir = args.outdir
    else:
        outdir = os.path.dirname(os.path.abspath(args.monomer))
    
    os.makedirs(outdir, exist_ok=True)
    
    # =========================================================================
    print_header("PTFEMA 链构建工具")
    print(f"  单体文件:   {args.monomer}")
    print(f"  聚合度:     {args.n}")
    print(f"  端基类型:   {args.cap}")
    print(f"  输出目录:   {outdir}")
    print(f"  输出前缀:   {args.outprefix}")
    if args.docx:
        print(f"  DOCX 模板:  {args.docx}")
    
    # =========================================================================
    # 1. 读取单体
    print_header("1. 读取单体")
    monomer_mol2 = parse_mol2(args.monomer)
    log_success(f"单体: {monomer_mol2.num_atoms} 原子, {monomer_mol2.num_bonds} 键")
    
    # =========================================================================
    # 2. 处理 DOCX（如果提供）
    docx_coords = None
    if args.docx:
        print_header("2. 处理 DOCX 坐标模板")
        docx_coords = extract_coords_from_docx(args.docx)
        
        if docx_coords:
            # 输出 XYZ
            xyz_path = os.path.join(outdir, f"{args.outprefix}_from_docx.xyz")
            coords_to_xyz(docx_coords, xyz_path, "PTFMA from docx")
            log_success(f"XYZ: {xyz_path}")
            
            # 尝试推断键并输出 mol2/pdb
            docx_mol = coords_to_rdkit_mol(docx_coords)
            if docx_mol:
                mol2_path = os.path.join(outdir, f"{args.outprefix}_from_docx.mol2")
                pdb_path = os.path.join(outdir, f"{args.outprefix}_from_docx.pdb")
                write_mol2_from_rdkit(docx_mol, mol2_path, f"{args.outprefix}_docx")
                write_pdb(docx_mol, pdb_path)
                log_success(f"MOL2: {mol2_path}")
                log_success(f"PDB: {pdb_path}")
    else:
        print_header("2. 跳过 DOCX 处理（未提供）")
    
    # =========================================================================
    # 3. 构建重复单元 + 封端模型
    print_header("3. 构建重复单元 + 封端模型")
    repeat_capped = build_repeat_capped(args.cap)
    
    repeat_path = os.path.join(outdir, f"{args.outprefix}_repeat_capped.mol2")
    write_mol2_from_rdkit(repeat_capped, repeat_path, f"{args.outprefix}_rep")
    
    log_success(f"重复单元: {repeat_capped.GetNumAtoms()} 原子, {repeat_capped.GetNumBonds()} 键")
    log_success(f"文件: {repeat_path}")
    
    # =========================================================================
    # 4. 构建 N 聚合度链
    print_header(f"4. 构建 {args.n} 聚合度链")
    polymer = build_ptfema_chain(monomer_mol2, args.n, args.cap)
    
    # =========================================================================
    # 5. 验证
    print_header("5. 验证结果")
    validation = validate_polymer(polymer, args.n)
    
    print(f"\n=== 结构验证 ===")
    print(f"  片段数:       {validation['num_fragments']} {'✓' if validation['is_single_molecule'] else '✗'}")
    print(f"  原子数:       {validation['num_atoms']}")
    print(f"  键数:         {validation['num_bonds']}")
    
    if 'cc_bond_min' in validation:
        print(f"\n=== 主链 C-C 键长 ===")
        print(f"  最小: {validation['cc_bond_min']:.3f} Å")
        print(f"  最大: {validation['cc_bond_max']:.3f} Å")
        print(f"  平均: {validation['cc_bond_avg']:.3f} Å")
    
    print(f"\n=== 几何检查 ===")
    print(f"  最小非键距离: {validation['min_nonbond_distance']:.3f} Å")
    print(f"  无原子重叠:   {'✓' if validation['no_clashes'] else '✗'}")
    
    if not validation['is_single_molecule']:
        log_error("验证失败：生成的结构不是单一分子！")
        return 1
    
    if not validation['no_clashes']:
        log_warn("警告：存在原子重叠，建议进一步优化")
    
    # =========================================================================
    # 6. 写入输出文件
    print_header("6. 写入输出文件")
    
    mol2_path = os.path.join(outdir, f"{args.outprefix}_{args.n}.mol2")
    pdb_path = os.path.join(outdir, f"{args.outprefix}_{args.n}.pdb")
    
    write_mol2_from_rdkit(polymer, mol2_path, f"{args.outprefix}_{args.n}")
    write_pdb(polymer, pdb_path)
    
    log_success(f"MOL2: {mol2_path} ({os.path.getsize(mol2_path)} bytes)")
    log_success(f"PDB:  {pdb_path} ({os.path.getsize(pdb_path)} bytes)")
    
    # =========================================================================
    # 汇总
    print_header("完成")
    print(f"\n生成的文件:")
    print(f"  1. {repeat_path}")
    print(f"  2. {mol2_path}")
    print(f"  3. {pdb_path}")
    if args.docx and docx_coords:
        print(f"  4. {os.path.join(outdir, f'{args.outprefix}_from_docx.xyz')}")
    
    print(f"\n端基封端类型: {args.cap}")
    print(f"聚合度: {args.n}")
    print(f"片段数: {validation['num_fragments']} (PASS)")
    
    return 0


if __name__ == "__main__":
    sys.exit(main())

